"""
Document Processor for extracting text from various file formats.

Supports:
- PDF files (using pdfplumber)
- DOCX files (using python-docx)
- Images (using pytesseract OCR)
"""

import os
import io
from pathlib import Path
from typing import List, Dict, Optional
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image


class DocumentProcessor:
    """Process documents and extract text with metadata."""
    
    def __init__(self):
        """Initialize document processor."""
        # Configure Tesseract path if needed (for Docker, it's in PATH)
        # For local development, you might need to set:
        # pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
        pass
    
    def process_file(self, file_path: str, file_content: Optional[bytes] = None) -> List[Dict[str, any]]:
        """
        Process a single file and extract text chunks with metadata.
        
        Args:
            file_path: Path to the file or filename
            file_content: Optional file content as bytes (for uploaded files)
        
        Returns:
            List of dictionaries with 'text' and 'metadata' keys
        """
        file_ext = Path(file_path).suffix.lower()
        filename = Path(file_path).name
        
        if file_ext == '.pdf':
            return self._process_pdf(file_path, filename, file_content)
        elif file_ext in ['.docx', '.doc']:
            return self._process_docx(file_path, filename, file_content)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']:
            return self._process_image(file_path, filename, file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _process_pdf(self, file_path: str, filename: str, file_content: Optional[bytes] = None) -> List[Dict[str, any]]:
        """Extract text from PDF file."""
        chunks = []
        
        try:
            if file_content:
                # Process from bytes (uploaded file)
                pdf_file = io.BytesIO(file_content)
                pdf = pdfplumber.open(pdf_file)
            else:
                # Process from file path
                pdf = pdfplumber.open(file_path)
            
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    chunks.append({
                        'text': text.strip(),
                        'metadata': {
                            'source': filename,
                            'page': page_num,
                            'file_type': 'pdf',
                            'total_pages': len(pdf.pages)
                        }
                    })
            
            pdf.close()
        except Exception as e:
            raise ValueError(f"Error processing PDF {filename}: {str(e)}")
        
        return chunks
    
    def _process_docx(self, file_path: str, filename: str, file_content: Optional[bytes] = None) -> List[Dict[str, any]]:
        """Extract text from DOCX file."""
        chunks = []
        
        try:
            if file_content:
                # Process from bytes (uploaded file)
                doc_file = io.BytesIO(file_content)
                doc = Document(doc_file)
            else:
                # Process from file path
                doc = Document(file_path)
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Combine paragraphs into chunks
            if paragraphs:
                full_text = '\n\n'.join(paragraphs)
                chunks.append({
                    'text': full_text,
                    'metadata': {
                        'source': filename,
                        'file_type': 'docx',
                        'paragraph_count': len(paragraphs)
                    }
                })
            
        except Exception as e:
            raise ValueError(f"Error processing DOCX {filename}: {str(e)}")
        
        return chunks
    
    def _process_image(self, file_path: str, filename: str, file_content: Optional[bytes] = None) -> List[Dict[str, any]]:
        """Extract text from image using OCR."""
        chunks = []
        
        try:
            if file_content:
                # Process from bytes (uploaded file)
                image = Image.open(io.BytesIO(file_content))
            else:
                # Process from file path
                image = Image.open(file_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            if text and text.strip():
                # Convert image.size tuple to string format (ChromaDB doesn't support tuples)
                image_size = image.size  # This is a tuple (width, height)
                chunks.append({
                    'text': text.strip(),
                    'metadata': {
                        'source': filename,
                        'file_type': 'image',
                        'image_width': image_size[0],  # Store width as int
                        'image_height': image_size[1],  # Store height as int
                        'image_size': f"{image_size[0]}x{image_size[1]}",  # Store as string for display
                        'image_mode': image.mode
                    }
                })
            
        except Exception as e:
            raise ValueError(f"Error processing image {filename}: {str(e)}")
        
        return chunks
    
    def process_multiple_files(
        self,
        files: List[tuple[str, bytes]]
    ) -> List[Dict[str, any]]:
        """
        Process multiple files.
        
        Args:
            files: List of tuples (filename, file_content_bytes)
        
        Returns:
            List of all text chunks from all files
        """
        all_chunks = []
        
        for filename, file_content in files:
            try:
                chunks = self.process_file(filename, file_content)
                all_chunks.extend(chunks)
            except Exception as e:
                # Log error but continue with other files
                print(f"Warning: Failed to process {filename}: {str(e)}")
                continue
        
        return all_chunks
